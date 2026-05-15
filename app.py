import streamlit as st
import requests
import io
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image
import pytesseract

# ================= 配置区 =================
API_KEY = st.secrets.get("ZHIPU_API_KEY", "你的API_KEY填这里")
API_URL = "https://open.bigmodel.cn/api/paas/v4/chat/completions"

st.set_page_config(page_title="智盾·天机教学提炼系统", layout="centered")

# ================= 沉浸式 UI =================
st.markdown("""
<meta name="google" content="notranslate">
<style>
    .stApp { background-color: #F8F9FA; color: #212529; }
    .block-container { padding: 2rem 1rem !important; }
    header { visibility: hidden !important; }
    footer { visibility: hidden !important; }
    .stButton button { 
        width: 100%; height: 50px; font-size: 16px; font-weight: bold;
        border-radius: 8px; background-color: #0056B3 !important;
        color: #FFFFFF !important; border: none; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .stTextArea textarea, .stTextInput input { 
        font-size: 15px !important; background-color: #FFFFFF !important;
        color: #212529 !important; border: 1px solid #CED4DA !important; border-radius: 8px;
    }
    [data-testid="stFileUploader"] {
        background-color: #FFFFFF; border: 1px dashed #0056B3; border-radius: 8px;
    }
    .question-card { background: #fff; padding: 15px; border-radius: 8px; border: 1px solid #e0e0e0; margin-bottom: 15px; box-shadow: 0 2px 4px rgba(0,0,0,0.05);}
</style>
""", unsafe_allow_html=True)

# 初始化 Session State
if "extracted_points" not in st.session_state:
    st.session_state.extracted_points = ""
if "quiz_data" not in st.session_state:
    st.session_state.quiz_data = None

# ================= 工具函数 =================
def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_image(file):
    img = Image.open(file)
    img.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
    return pytesseract.image_to_string(img, lang='chi_sim+eng')

def call_glm_api(system_prompt, user_content):
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "glm-4-flash",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ],
        "temperature": 0.3, "max_tokens": 4096 
    }
    try:
        response = requests.post(API_URL, headers=headers, json=data)
        return response.json()['choices'][0]['message']['content']
    except Exception as e:
        return f"API异常: {e}"

def call_glm_api_stream(system_prompt, user_content):
    """流式调用接口，用于实现一个字一个字蹦出来的效果"""
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    data = {
        "model": "glm-4-flash",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_content}
        ],
        "temperature": 0.3, "max_tokens": 8192,
        "stream": True 
    }
    try:
        response = requests.post(API_URL, headers=headers, json=data, stream=True)
        for line in response.iter_lines():
            if line:
                decoded_line = line.decode('utf-8')
                if decoded_line.startswith("data: "):
                    content = decoded_line[6:]
                    if content == "[DONE]": break
                    chunk = json.loads(content)
                    yield chunk['choices'][0]['delta'].get('content', '')
    except Exception as e:
        yield f"\n[网络流异常: {e}]"

def extract_json_from_text(text):
    """最安全的JSON提取器"""
    try:
        start_idx = text.find('{')
        end_idx = text.rfind('}')
        if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
            clean_text = text[start_idx:end_idx+1]
            return json.loads(clean_text)
        return json.loads(text)
    except:
        return None

# ================= Word 生成核心引擎 =================
def set_chinese_font(run, font_name='宋体', size=12):
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def generate_professional_word(exam_data, version="teacher"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(exam_data.get('title', '智能生成专项练习卷'))
    title_run.bold = True
    set_chinese_font(title_run, '黑体', 18)

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run("姓名：___________  班级：___________  考号：___________  得分：___________")
    set_chinese_font(info_run, '楷体', 12)
    info_para.paragraph_format.space_after = Pt(20)

    questions = exam_data.get('questions', [])
    if not questions:
        doc.add_paragraph("未生成题目数据。")
        
    for idx, q_item in enumerate(questions):
        q_type = q_item.get('type', '未知题型')
        score = q_item.get('score', 0)
        
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{idx + 1}. [{q_type} - {score}分] {q_item.get('q', '')}")
        set_chinese_font(q_run, '宋体', 12)
        q_run.bold = True
        q_para.paragraph_format.space_after = Pt(5)

        options = q_item.get('options', [])
        if options and isinstance(options, list):
            for opt in options:
                if str(opt).strip():
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(0.75)
                    opt_para.paragraph_format.space_after = Pt(2)
                    opt_run = opt_para.add_run(str(opt))
                    set_chinese_font(opt_run, '宋体', 11)
        
        if q_type in ["简答题", "解答题"] and not options:
            doc.add_paragraph("\n\n")

        doc.add_paragraph().paragraph_format.space_after = Pt(5)

    if version == "teacher":
        doc.add_page_break()
        ans_title = doc.add_paragraph()
        ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ans_run = ans_title.add_run("参考答案与深度解析（教师专享）")
        ans_run.bold = True
        set_chinese_font(ans_run, '黑体', 16)
        
        for idx, q_item in enumerate(questions):
            ans_para = doc.add_paragraph()
            num_run = ans_para.add_run(f"【第{idx + 1}题】 ")
            set_chinese_font(num_run, '黑体', 12)
            
            ans_label = ans_para.add_run(f"正确答案：{q_item.get('answer', '略')}\n")
            ans_label.font.color.rgb = RGBColor(255, 0, 0)
            set_chinese_font(ans_label, '宋体', 12)
            ans_label.bold = True
            
            analysis_run = ans_para.add_run(f"详细解析：{q_item.get('analysis', '略')}")
            set_chinese_font(analysis_run, '楷体', 11)
            ans_para.paragraph_format.space_after = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ================= 主界面 =================
st.title("智盾·天机 教学提炼与出题系统")

# ================= 步骤 1：素材摄入 =================
st.header("步骤1：教学素材深度剖析")
col1, col2 = st.columns([2, 1])

with col1:
    tab1, tab2 = st.tabs(["📄 文件上传", "⌨️ 手动粘贴"])
    with tab1:
        uploaded_files = st.file_uploader("上传 Word 课件或讲义照片", type=['docx', 'png', 'jpg', 'jpeg'], accept_multiple_files=True)
        parsed_text = ""
        if uploaded_files:
            with st.spinner(f"正在解析 {len(uploaded_files)} 个文件..."):
                for file in uploaded_files:
                    if file.name.endswith('.docx'):
                        parsed_text += extract_text_from_docx(file) + "\n\n"
                    else:
                        parsed_text += extract_text_from_image(file) + "\n\n"
            if len(parsed_text.strip()) > 0:
                st.success(f"解析成功：共 {len(parsed_text)} 字符")
                
    with tab2:
        manual_text = st.text_area("此处输入文本：", height=150)
    
    final_source = parsed_text if uploaded_files else manual_text

with col2:
    st.info(" **系统提示**：AI将智能剥离冗余信息，提取核心考点。")
    if st.button("一键提取核心重难点", type="primary"):
        if len(final_source.strip()) < 15:
            st.error("有效文本不足 15 字符。")
        else:
            with st.spinner("AI 正在萃取知识骨架..."):
                sys_prompt_step1 = """你是一名特级教师。提取核心概念、教学难点、考点预测。必须输出纯文本格式，禁用Markdown符号。格式如：\n一、核心概念\n1. 内容\n二、教学难点..."""
                raw = call_glm_api(sys_prompt_step1, final_source)
                st.session_state.extracted_points = raw.replace("*", "").replace("#", "")
                st.rerun()

st.divider()

# ================= 步骤 2：动态组卷 =================
st.header("步骤2：混合题型组卷配置")
edited_points = st.text_area("确认知识图谱（可修改）：", value=st.session_state.extracted_points, height=200)

st.markdown("#####  定制试卷结构")
c_type1, c_type2, c_type3 = st.columns(3)
with c_type1:
    count_single = st.number_input("单选题数量", min_value=0, max_value=20, value=3)
    score_single = st.number_input("单选题单题分值", min_value=1, max_value=10, value=5)
with c_type2:
    count_fill = st.number_input("填空题数量", min_value=0, max_value=20, value=2)
    score_fill = st.number_input("填空题单题分值", min_value=1, max_value=10, value=5)
with c_type3:
    count_short = st.number_input("简答题数量", min_value=0, max_value=10, value=1)
    score_short = st.number_input("简答题单题分值", min_value=1, max_value=30, value=15)

paper_title = st.text_input("试卷大标题", "智能生成课后专项练习卷")

if st.button("开始生成结构化试卷", type="primary"):
    total_q = count_single + count_fill + count_short
    if total_q == 0:
        st.warning("请至少选择一道题目生成！")
    else:
        sys_prompt_step2 = f"""你是一个严谨的教育命题专家。根据知识点生成试卷。
必须且只能输出合法的 JSON 格式数据。不要有任何额外的解释文字，绝对不要用任何符号包裹数据。
你需要生成：
- 单选题: {count_single}道 (每题{score_single}分)
- 填空题: {count_fill}道 (每题{score_fill}分)
- 简答题: {count_short}道 (每题{score_short}分)

JSON 结构严格遵守以下格式：
{{
  "title": "{paper_title}",
  "questions": [
    {{
      "type": "单项选择题",
      "score": {score_single},
      "q": "题干内容...",
      "options": ["A. 选项1", "B. 选项2", "C. 选项3", "D. 选项4"],
      "answer": "A",
      "analysis": "考点解析..."
    }},
    {{
      "type": "填空题",
      "score": {score_fill},
      "q": "题干内容，空白处用横线______表示...",
      "options": [],
      "answer": "填空答案",
      "analysis": "考点解析..."
    }},
    {{
      "type": "简答题",
      "score": {score_short},
      "q": "简答题题干...",
      "options": [],
      "answer": "详细的简答题标准答案要点...",
      "analysis": "考点解析..."
    }}
  ]
}}"""
        
        st.subheader("正在构思题目并构建数据...")
        
        # 【新增：打字机流式输出容器】
        stream_container = st.empty()
        
        with stream_container:
            # 这里调用打字机效果，把 AI 生成底层代码的过程展现给用户看
            stream_gen = call_glm_api_stream(sys_prompt_step2, edited_points)
            raw_json_str = st.write_stream(stream_gen)
            
        # 打字结束，瞬间进行解析转换
        parsed_data = extract_json_from_text(raw_json_str)
        
        if parsed_data and "questions" in parsed_data:
            st.session_state.quiz_data = parsed_data
            # 清空刚才打字的原始代码，准备展示漂亮的 UI
            stream_container.empty()
            st.rerun()  # 触发刷新，立刻显示步骤3
        else:
            st.error("解析失败，大模型返回格式异常。请重试。")

# ================= 步骤 3 & 4：微调与导出 =================
if st.session_state.quiz_data:
    st.divider()
    st.header("步骤3：试卷在线预览与微调 (最核心)")
    st.info(" **试卷生成完毕！** 您可以直接在下方修改 AI 生成的题目，修改后将自动同步到最终导出的 Word 中。")
    
    quiz = st.session_state.quiz_data
    quiz['title'] = st.text_input(" 试卷主标题", quiz.get('title', ''))

    for i, q in enumerate(quiz.get('questions', [])):
        with st.expander(f"第 {i+1} 题 - {q.get('type', '')} ({q.get('score', 0)}分)", expanded=False):
            q['q'] = st.text_area("题干", q.get('q', ''), key=f"q_{i}", height=80)
            
            if q.get('options'):
                opts_text = "\n".join(q['options'])
                new_opts = st.text_area("选项 (每行一个)", opts_text, key=f"opt_{i}", height=100)
                q['options'] = [o.strip() for o in new_opts.split('\n') if o.strip()]
            
            col_a, col_b = st.columns(2)
            with col_a:
                q['answer'] = st.text_input("正确答案", q.get('answer', ''), key=f"ans_{i}")
            with col_b:
                q['analysis'] = st.text_area("深度解析", q.get('analysis', ''), key=f"aly_{i}", height=100)

    st.divider()
    st.header("步骤4：双版本精美排版导出")
    st.markdown(" **您可以将微调后的试卷导出为标准的教学 Word 文档。**")
    
    col_dl1, col_dl2 = st.columns(2)
    
    with col_dl1:
        student_doc = generate_professional_word(st.session_state.quiz_data, version="student")
        st.download_button(
            label=" 下载【学生版】试卷 (纯题目版)",
            data=student_doc,
            file_name=f"{quiz['title']}_学生版.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        st.caption("适用场景：直接打印下发给学生测试")
        
    with col_dl2:
        teacher_doc = generate_professional_word(st.session_state.quiz_data, version="teacher")
        st.download_button(
            label=" 下载【教师版】试卷 (含答案解析)",
            data=teacher_doc,
            file_name=f"{quiz['title']}_教师解析版.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        st.caption("适用场景：教师备课、课后批改讲解")
